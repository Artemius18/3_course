import docx

def encode_message(original_doc_path, message, output_doc_path):
    doc = docx.Document(original_doc_path)
    message += ' '  # Добавляем пробел в конец сообщения для обозначения конца
    message_bin = ''.join(format(ord(i), '08b') for i in message)  # 01110011 01100101 01100011 01110010 01100101 01110100 + 00100000 (это пробел выше)
    # print(message_bin)
    for paragraph in doc.paragraphs:
        words = paragraph.text.split()
        # print(words)
        encoded_words = []

        for word in words:
            if message_bin:
                print(message_bin)
                bit = message_bin[0]
                print('bit: ', bit)
                message_bin = message_bin[1:]

                if bit == '1':
                    encoded_words.append(word + '  ')  # Добавляем двойной пробел после слова
                    print(encoded_words)
                else:
                    encoded_words.append(word + ' ')  # Добавляем одинарный пробел после слова
                    print(encoded_words)
            else:
                encoded_words.append(word + ' ')  # Если сообщение закончилось, добавляем одинарный пробел

        paragraph.text = ''.join(encoded_words)

    doc.save(output_doc_path)

def decode_message(encoded_doc_path):
    doc = docx.Document(encoded_doc_path)
    message_bin = ''

    for paragraph in doc.paragraphs:
        words = paragraph.text.split(' ')
        for i in range(len(words) - 1):  # Исключаем последнее слово, так как оно может не содержать пробела
            if words[i+1] == '':  # Если следующее слово пустое, это был двойной пробел
                message_bin += '1'
            elif words[i]:  # Если текущее слово не пустое, это был одинарный пробел
                message_bin += '0'

    message = ''
    for i in range(0, len(message_bin), 8):  # Разбиваем двоичное сообщение на байты
        byte = message_bin[i:i+8]
        message += chr(int(byte, 2))  # Преобразуем каждый байт в символ

    return message.rstrip()  # Удаляем пробел в конце сообщения

encode_message('original.docx', 'secret', 'encoded.docx')
message = decode_message('encoded.docx')
print(message)
